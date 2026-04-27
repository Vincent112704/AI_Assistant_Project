from docx import Document
from pydantic import BaseModel

class Sections(BaseModel):
    section_heading: str
    section_content: str

class DocumentContent(BaseModel):
    sections: list[Sections]


    



def document_generator(title: str, content: DocumentContent) -> Document:
    """
    A tool that creates a docx file with a given title and context
    
    Args:
        title: The title of the document
        content: [
            { 
                section_heading: "Intoduction",
                section_content: "This is the introduction section of the document" 
            }
        ]
    Returns:
        a Document object that can be saved as a docx file
    """

    docx = Document()
    docx.add_heading(title, level=0)

    for section in content:
        docx.add_heading(section.section_heading, level=1)
        docx.add_paragraph(section.section_content)
    
    return docx



title = "AI Adoption Report"

content = DocumentContent(
    sections=[
        Sections(
            section_heading="Introduction",
            section_content="This is the introduction section of the document."
        ),
        Sections(
            section_heading="Benefits of AI",
            section_content="AI helps businesses automate tasks and improve efficiency."
        ),
        Sections(
            section_heading="Conclusion",
            section_content="AI adoption will continue to grow across industries."
        )
    ]
)

doc = document_generator(title, content.sections)
doc.save("test.docx")





    
